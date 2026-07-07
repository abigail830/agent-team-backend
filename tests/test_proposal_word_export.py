import uuid
import re
import zipfile
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from app.proposal.draft import materialize_draft
from app.proposal.export_service import ProposalExportError, generate_proposal_docx, word_export_status
from app.proposal.storage import resolve_artifact_path
from app.proposal.word_context import build_word_context, cover_for_name, word_export_filename
from app.proposal.word_render import render_word_document


def test_build_word_context_payment_options_au_advisory():
    from tests.proposal_fee_fixtures import make_mdm_fee_row
    from app.proposal.draft import enable_draft_section

    draft = materialize_draft(template_id="au-advisory")
    fee = next(s for s in draft["document"]["sections"] if s["kind"] == "fee_section")
    fee["tables"] = [
        {
            "id": "table_setup",
            "title": "Setup of Xero",
            "rows": [
                make_mdm_fee_row(
                    {
                        "sku": "XERO",
                        "service_name": "Setup of Xero",
                        "scope_of_work": "",
                        "price_amount": 500.0,
                        "price_currency": "AUD",
                        "billing_frequency": "ONE_TIME",
                        "recurring": "ONE_OFF",
                        "pricing_type": "FIXED",
                    }
                )
            ],
        },
        {
            "id": "table_tax",
            "title": "Tax Package 2",
            "rows": [
                make_mdm_fee_row(
                    {
                        "sku": "TAX",
                        "service_name": "Application",
                        "scope_of_work": "",
                        "price_amount": 400.0,
                        "price_currency": "AUD",
                        "billing_frequency": "MONTHLY",
                        "recurring": "RECURRING",
                        "pricing_type": "FIXED",
                    }
                )
            ],
        },
    ]
    draft = enable_draft_section(draft, "payment_options")
    ctx = build_word_context(draft)
    payment = ctx["payment_options"]
    assert payment["has_options"] is True
    assert payment["currency"] == "AUD"
    assert len(payment["options"]) == 1
    option = payment["options"][0]
    assert option["label"] == "Payment Option A"
    assert len(option["rows"]) == 2
    assert option["rows"][0]["label"] == "Setup of Xero"
    assert option["rows"][0]["once_off_display"] == "AUD $500.00"
    assert option["rows"][1]["monthly_display"] == "AUD $400.00"
    assert option["rows"][1]["total_display"] == "AUD $4,800.00"
    assert option["summary"]["once_off_total_display"] == "AUD $500.00"
    assert option["summary"]["recurring_annualized_total_display"] == "AUD $4,800.00"


def test_word_export_without_optional_sections_in_draft():
    """Word templates use sections.<id>; missing draft nodes must get template stubs."""
    from jinja2 import Environment

    from app.proposal.loaders import load_template_yaml

    env = Environment(autoescape=False)
    cases = {
        "sg-incorp": ["appendices", "first_invoice"],
        "au-advisory": ["appendices", "payment_options", "credentials"],
        "harneys-bvi": ["appendices", "required_documents", "additional_info"],
    }

    for template_id, optional_ids in cases.items():
        draft = materialize_draft(template_id=template_id)
        draft["document"]["sections"] = [
            section
            for section in draft["document"]["sections"]
            if section.get("id") not in optional_ids
        ]
        ctx = build_word_context(draft)
        tpl = load_template_yaml(template_id)
        for spec in tpl.get("sections") or []:
            section_id = str(spec.get("id") or "").strip()
            if not section_id:
                continue
            assert section_id in ctx["sections"], template_id
            if section_id in optional_ids:
                assert ctx["sections"][section_id].enabled is False, section_id

        jinja_checks = [
            "{% if sections.appendices.enabled and sections.appendices.items %}{% endif %}",
        ]
        if template_id == "sg-incorp":
            jinja_checks.append(
                "{% if sections.first_invoice.enabled and first_invoice.has_rows %}{% endif %}"
            )
        if template_id == "au-advisory":
            jinja_checks.append(
                "{% if sections.payment_options.enabled and payment_options.has_options %}{% endif %}"
            )
        if template_id == "harneys-bvi":
            jinja_checks.append(
                "{% if sections.required_documents.enabled and sections.required_documents.has_content %}{% endif %}"
            )
        for template in jinja_checks:
            env.from_string(template).render(**ctx)


def test_word_export_without_appendices_section_in_draft():
    from pathlib import Path

    draft = materialize_draft(template_id="harneys-bvi")
    draft["document"]["sections"] = [
        section for section in draft["document"]["sections"] if section.get("id") != "appendices"
    ]
    ctx = build_word_context(draft)
    appendices = ctx["sections"]["appendices"]
    assert appendices.enabled is False
    assert appendices.items == []

    path = Path("agents/proposal-composer/knowledge/templates/harneys-bvi/export/proposal.docx")
    output = render_word_document(path, ctx)
    assert output.startswith(b"PK")


def test_build_word_context_bvi_aggregate_footnotes():
    from tests.proposal_fee_fixtures import make_mdm_fee_row

    draft = materialize_draft(template_id="harneys-bvi")
    fee_idx = next(
        i for i, s in enumerate(draft["document"]["sections"]) if s["kind"] == "fee_section"
    )
    draft["document"]["sections"][fee_idx]["tables"] = [
        {
            "id": "t1",
            "title": "Fees",
            "rows": [
                make_mdm_fee_row(
                    {
                        "sku": "CSS013",
                        "service_name": "Financial Annual Return (FAR)",
                        "description": "FAR filing",
                        "price_amount": 350.0,
                        "price_currency": "USD",
                        "billing_frequency": "ANNUALLY",
                        "pricing_type": "FIXED",
                        "footnotes": "Includes collecting and filing the completed FAR.",
                    },
                    template_id="harneys-bvi",
                ),
            ],
        }
    ]
    ctx = build_word_context(draft)
    footnotes = ctx["fee_tables"]["footnotes"]
    assert len(footnotes) == 1
    assert footnotes[0]["number"] == 1
    assert "FAR" in footnotes[0]["text"]


def test_cover_for_name_prefers_company():
    assert cover_for_name({"company_name": "Acme Pte Ltd", "contract_name": "Sara"}) == "Acme Pte Ltd"


def test_cover_for_name_falls_back_to_contact():
    assert cover_for_name({"contract_name": "Sara"}) == "Sara"
    assert cover_for_name({}) == ""


def test_richtext_filter_resolves_section_objects():
    from app.proposal.word_context import WordSectionContext, WordSectionIntro, resolve_word_text

    section = WordSectionContext(
        enabled=True,
        title="Scope",
        body="**Hello** world",
    )
    intro = WordSectionIntro("Intro line")
    assert resolve_word_text(section) == section.plain
    assert resolve_word_text(intro) == intro.plain
    assert str(section) == section.plain
    assert str(intro) == intro.plain


def test_appendix_blocks_page_break_after():

    draft = {
        "meta": {"template_id": "sg-incorp"},
        "facts": {"client": {}},
        "document": {
            "sections": [
                {
                    "id": "appendices",
                    "kind": "collection",
                    "title": "Appendices",
                    "enabled": True,
                    "blocks": [
                        {
                            "id": "a1",
                            "kind": "markdown_block",
                            "title": "Appendix A",
                            "enabled": True,
                            "content": "Body A",
                        },
                        {
                            "id": "a2",
                            "kind": "markdown_block",
                            "title": "Appendix B",
                            "enabled": True,
                            "content": "Body B",
                        },
                    ],
                }
            ]
        },
    }
    ctx = build_word_context(draft)
    blocks = ctx["sections"]["appendices"].blocks
    assert len(blocks) == 2
    assert ctx["sections"]["appendices"].items is blocks
    assert blocks[0].page_break_after is True
    assert blocks[1].page_break_after is False


def test_build_word_context_cover_for(tmp_path, monkeypatch):
    from app.proposal import loaders

    template_root = tmp_path / "sg-incorp"
    export_dir = template_root / "export"
    export_dir.mkdir(parents=True)
    doc = Document()
    doc.add_paragraph("For {{ cover_for }}")
    doc.save(export_dir / "proposal.docx")

    (template_root / "template.yaml").write_text(
        """
template_id: sg-incorp
display_name: Singapore Incorp
sections:
  - id: about_incorp
    kind: static_block
    title: About
    source:
      type: template_file
      file: blocks/about.md
document_export:
  word:
    enabled: true
    template_file: export/proposal.docx
""".strip(),
        encoding="utf-8",
    )
    blocks_dir = template_root / "blocks"
    blocks_dir.mkdir(parents=True)
    (blocks_dir / "about.md").write_text("About us", encoding="utf-8")

    monkeypatch.setattr(loaders, "TEMPLATES_ROOT", tmp_path)
    loaders.load_template_yaml.cache_clear()

    draft = materialize_draft(template_id="sg-incorp")
    draft["facts"]["client"]["contract_name"] = "Sara"
    ctx = build_word_context(draft)
    assert ctx["cover_for"] == "Sara"

    draft["facts"]["client"]["company_name"] = "Acme Pte Ltd"
    ctx = build_word_context(draft)
    assert ctx["cover_for"] == "Acme Pte Ltd"


def test_render_word_document_cover_for(tmp_path, monkeypatch):
    from app.proposal import loaders

    template_root = tmp_path / "sg-incorp"
    export_dir = template_root / "export"
    export_dir.mkdir(parents=True)
    template_path = export_dir / "proposal.docx"
    doc = Document()
    doc.add_paragraph("For {{ cover_for }}")
    doc.save(template_path)

    monkeypatch.setattr(loaders, "TEMPLATES_ROOT", tmp_path)

    output = render_word_document(template_path, {"cover_for": "Acme Pte Ltd"})
    assert output.startswith(b"PK")
    assert len(output) > 100


def test_generate_proposal_docx_persists_artifact(tmp_path, monkeypatch):
    from app.proposal import loaders

    template_root = tmp_path / "sg-incorp"
    export_dir = template_root / "export"
    export_dir.mkdir(parents=True)
    doc = Document()
    doc.add_paragraph("For {{ cover_for }}")
    doc.save(export_dir / "proposal.docx")

    (template_root / "template.yaml").write_text(
        """
template_id: sg-incorp
display_name: Singapore Incorp
sections:
  - id: about_incorp
    kind: static_block
    title: About
    source:
      type: template_file
      file: blocks/about.md
document_export:
  word:
    enabled: true
    template_file: export/proposal.docx
""".strip(),
        encoding="utf-8",
    )
    blocks_dir = template_root / "blocks"
    blocks_dir.mkdir(parents=True)
    (blocks_dir / "about.md").write_text("About us", encoding="utf-8")

    monkeypatch.setattr(loaders, "TEMPLATES_ROOT", tmp_path)
    loaders.load_template_yaml.cache_clear()

    draft = materialize_draft(template_id="sg-incorp")
    draft["facts"]["client"]["company_name"] = "Acme Pte Ltd"
    from app.proposal.placeholders import sync_draft_template_placeholders

    draft = sync_draft_template_placeholders(draft)
    chat_id = uuid.uuid4()

    result = generate_proposal_docx(draft, chat_id=chat_id, force=False, persist=True)
    assert result["status"] == "ok"
    assert result["filename"].endswith(".docx")
    assert result["download_url"]

    path = resolve_artifact_path(chat_id, result["artifact_id"])
    assert path is not None
    assert path.suffix == ".docx"
    assert path.read_bytes().startswith(b"PK")


def test_word_export_status_without_template_file(tmp_path, monkeypatch):
    from app.proposal import loaders

    template_root = tmp_path / "sg-incorp"
    template_root.mkdir(parents=True)
    (template_root / "template.yaml").write_text("template_id: sg-incorp\nsections: []\n", encoding="utf-8")
    monkeypatch.setattr(loaders, "TEMPLATES_ROOT", tmp_path)
    loaders.load_template_yaml.cache_clear()

    draft = {"meta": {"template_id": "sg-incorp"}, "facts": {"client": {}}}
    status = word_export_status(draft)
    assert status["available"] is False
    assert status["reason"] == "no_word_template"


def test_generate_proposal_docx_blocked_without_content():
    draft = {"meta": {"template_id": "sg-incorp"}, "facts": {"client": {}}, "document": {"sections": []}}
    with pytest.raises(ProposalExportError) as exc:
        generate_proposal_docx(draft, force=False, persist=False)
    assert exc.value.code == "blocked"


def test_word_export_filename_uses_meta_title_and_cover_for():
    draft = {
        "meta": {"title": "FEE PROPOSAL - Oversee Limited"},
        "facts": {"client": {"company_name": "Oversee Limited", "contract_name": "Ably"}},
    }
    assert word_export_filename(draft) == "FEE PROPOSAL - Oversee Limited.docx"


def test_word_export_filename_appends_cover_when_not_in_title():
    draft = {
        "meta": {"title": "FEE PROPOSAL - Sara"},
        "facts": {"client": {"company_name": "Oversee Limited", "contract_name": "Sara"}},
    }
    assert word_export_filename(draft) == "FEE PROPOSAL - Sara - Oversee Limited.docx"


def test_render_word_document_deduplicates_word_ids(tmp_path, monkeypatch):
    from collections import Counter

    from app.proposal import loaders

    template_root = tmp_path / "sg-incorp"
    export_dir = template_root / "export"
    export_dir.mkdir(parents=True)
    template_path = export_dir / "proposal.docx"
    doc = Document()
    p = doc.add_paragraph("Line 1")
    p2 = doc.add_paragraph("Line 2")
    p._element.set(
        "{http://schemas.microsoft.com/office/word/2010/wordml}paraId",
        "A1B2C3D4",
    )
    p._element.set(
        "{http://schemas.microsoft.com/office/word/2010/wordml}textId",
        "11111111",
    )
    p2._element.set(
        "{http://schemas.microsoft.com/office/word/2010/wordml}paraId",
        "A1B2C3D4",
    )
    p2._element.set(
        "{http://schemas.microsoft.com/office/word/2010/wordml}textId",
        "11111111",
    )
    doc.save(template_path)

    (template_root / "template.yaml").write_text(
        "template_id: sg-incorp\nsections: []\n",
        encoding="utf-8",
    )
    monkeypatch.setattr(loaders, "TEMPLATES_ROOT", tmp_path)

    output = render_word_document(template_path, {"cover_for": "Acme"})
    with zipfile.ZipFile(BytesIO(output)) as zout:
        xml = zout.read("word/document.xml").decode("utf-8")
    para_ids = re.findall(r'w14:paraId="([^"]+)"', xml)
    text_ids = re.findall(r'w14:textId="([^"]+)"', xml)
    assert all(count == 1 for count in Counter(para_ids).values())
    assert all(count == 1 for count in Counter(text_ids).values())
