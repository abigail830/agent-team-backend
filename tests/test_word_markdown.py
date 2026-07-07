from io import BytesIO
import re
import zipfile
from pathlib import Path

import pytest
from docx import Document

from app.proposal.draft import materialize_draft
from app.proposal.word_context import build_word_context
from app.proposal.word_markdown import (
    build_markdown_subdoc,
    markdown_has_gfm_table,
    markdown_to_plain,
    split_gfm_markdown,
)
from app.proposal.word_render import render_word_document
from docxtpl import DocxTemplate


FAR_APPENDIX_MARKDOWN = """## Financial Annual Return

Every BVI entity is required to provide financial information annually.

| Packages | Annual Fee (US$) | FAR storage |
| --- | --- | --- |
| Lite | $217 | ✓ |
| Standard | $630 | ✓ |

Footnote text after the table.
"""


def test_split_gfm_markdown_extracts_table_block():
    blocks = split_gfm_markdown(FAR_APPENDIX_MARKDOWN)
    assert blocks[0]["type"] == "text"
    assert "Financial Annual Return" in blocks[0]["raw"]
    assert blocks[1]["type"] == "table"
    assert blocks[1]["headers"] == ["Packages", "Annual Fee (US$)", "FAR storage"]
    assert blocks[1]["rows"][0] == ["Lite", "$217", "✓"]
    assert blocks[1]["rows"][1] == ["Standard", "$630", "✓"]
    assert blocks[-1]["type"] == "text"
    assert "Footnote text" in blocks[-1]["raw"]


def test_markdown_has_gfm_table():
    assert markdown_has_gfm_table(FAR_APPENDIX_MARKDOWN) is True
    assert markdown_has_gfm_table("Plain paragraph only.") is False


def test_markdown_to_plain_still_strips_headings():
    plain = markdown_to_plain("## Title\n\nBody line.")
    assert "Title" in plain
    assert "##" not in plain


def test_split_single_column_gfm_table():
    kyc = Path(
        "agents/proposal-composer/knowledge/peripheral/required-docs/harneys/KYC closing statement.md"
    ).read_text(encoding="utf-8")
    assert markdown_has_gfm_table(kyc) is True
    blocks = split_gfm_markdown(kyc)
    table = next(b for b in blocks if b["type"] == "table")
    assert table["headers"] == ["Content"]
    assert len(table["rows"]) == 8
    assert "Certified true copy" in table["rows"][0][0]


def test_build_markdown_subdoc_creates_word_table(tmp_path):
    template_path = tmp_path / "subdoc.docx"
    doc = Document()
    doc.add_paragraph("Before {{ appendix.subdoc }} After")
    doc.save(template_path)

    tpl = DocxTemplate(str(template_path))
    appendix = type("Block", (), {"body": FAR_APPENDIX_MARKDOWN, "subdoc": None})()
    appendix.subdoc = build_markdown_subdoc(tpl, appendix.body)

    output = render_word_document(
        template_path,
        {"appendix": appendix},
    )
    with zipfile.ZipFile(BytesIO(output)) as zout:
        xml = zout.read("word/document.xml").decode("utf-8")
    assert "<w:tbl" in xml
    assert "Lite" in xml
    assert "$217" in xml
    assert "✓" in xml
    assert "| Packages |" not in xml
    assert "Footnote text after the table." in xml


def test_build_word_context_attaches_subdoc_for_appendix_block():
    draft = materialize_draft(template_id="harneys-bvi")
    appendices = next(
        section for section in draft["document"]["sections"] if section["id"] == "appendices"
    )
    appendices["enabled"] = True
    appendices["blocks"] = [
        {
            "id": "far",
            "kind": "markdown_block",
            "title": "Appendix: Financial Annual Return (FAR)",
            "enabled": True,
            "content": FAR_APPENDIX_MARKDOWN,
        }
    ]
    ctx = build_word_context(draft)
    block = ctx["sections"]["appendices"].items[0]
    assert block.has_content is True
    assert markdown_has_gfm_table(block.body) is True
    assert block.subdoc is None

    template_path = Path("agents/proposal-composer/knowledge/templates/sg-incorp/export/proposal.docx")
    if not template_path.is_file():
        pytest.skip("sg-incorp proposal.docx template not available")
    tpl = DocxTemplate(str(template_path))
    block.subdoc = build_markdown_subdoc(tpl, block.body)
    assert block.subdoc is not None


def test_render_harneys_required_documents_kyc_table(tmp_path):
    from app.proposal import loaders

    template_root = tmp_path / "harneys-bvi"
    export_dir = template_root / "export"
    export_dir.mkdir(parents=True)
    src = Path("agents/proposal-composer/knowledge/templates/harneys-bvi/export/proposal.docx")
    if not src.is_file():
        pytest.skip("harneys-bvi proposal.docx not available")
    template_path = export_dir / "proposal.docx"
    template_path.write_bytes(src.read_bytes())
    (template_root / "template.yaml").write_text(
        "template_id: harneys-bvi\nsections: []\n",
        encoding="utf-8",
    )
    monkeypatch = pytest.MonkeyPatch()
    monkeypatch.setattr(loaders, "TEMPLATES_ROOT", tmp_path)

    kyc = Path(
        "agents/proposal-composer/knowledge/peripheral/required-docs/harneys/KYC closing statement.md"
    ).read_text(encoding="utf-8")
    draft = materialize_draft(template_id="harneys-bvi")
    req = next(s for s in draft["document"]["sections"] if s["id"] == "required_documents")
    req["content"] = f"## KYC closing statement\n\n{kyc}"

    ctx = build_word_context(draft)
    assert ctx["sections"]["required_documents"].subdoc is None
    output = render_word_document(template_path, ctx)
    with zipfile.ZipFile(BytesIO(output)) as zout:
        xml = zout.read("word/document.xml").decode("utf-8")
    assert "<w:tbl" in xml
    assert "Certified true copy" in xml
    assert "| Content |" not in xml
    monkeypatch.undo()


def test_render_word_document_includes_native_table_for_appendix(tmp_path, monkeypatch):
    from app.proposal import loaders

    template_root = tmp_path / "harneys-bvi"
    export_dir = template_root / "export"
    export_dir.mkdir(parents=True)
    template_path = export_dir / "proposal.docx"
    doc = Document()
    doc.add_paragraph(
        "{% for item in sections.appendices.items %}"
        "{{ item.title }}"
        "{{ item.subdoc }}"
        "{% endfor %}"
    )
    doc.save(template_path)
    (template_root / "template.yaml").write_text(
        "template_id: harneys-bvi\nsections: []\n",
        encoding="utf-8",
    )
    monkeypatch.setattr(loaders, "TEMPLATES_ROOT", tmp_path)

    draft = materialize_draft(template_id="harneys-bvi")
    appendices = next(
        section for section in draft["document"]["sections"] if section["id"] == "appendices"
    )
    appendices["enabled"] = True
    appendices["blocks"] = [
        {
            "id": "far",
            "kind": "markdown_block",
            "title": "Appendix: FAR",
            "enabled": True,
            "content": FAR_APPENDIX_MARKDOWN,
        }
    ]
    ctx = build_word_context(draft)
    output = render_word_document(template_path, ctx)
    with zipfile.ZipFile(BytesIO(output)) as zout:
        xml = zout.read("word/document.xml").decode("utf-8")
    assert "<w:tbl" in xml
    assert "Lite" in xml
    assert re.search(r"\$217", xml)
    assert "| Packages |" not in xml
